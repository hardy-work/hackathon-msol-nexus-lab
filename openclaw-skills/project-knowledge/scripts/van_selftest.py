#!/usr/bin/env python3
"""CI fixture cho LUỒNG VĂN: .docx/.pdf → raw → cổng Stage 4 → Gate 3a → Gate 4.

Đây từng là lane duy nhất không có selftest, và cũng là lane duy nhất chưa từng chạy
(`extract/van-docs.yml` rỗng, `run_all.sh` không gọi `ingest_van.py`). Hai lỗi đi kèm:

  1. Cổng Stage 4 chặn ĐÚNG cái mà prompt của chính nó bắt LLM viết. `check_transform`
     đọc đơn vị bằng chữ đứng ngay sau con số, nên trong `facts: 8, unit: "ký tự"` nó
     thấy cặp (8, không-đơn-vị) còn nguồn ghi (8, ký tự) → báo "số mới/đổi/làm tròn".
     Mọi trang khai số đều bị vứt; muốn trang được ghi thì phải bỏ trống `facts:`, và
     khi đó Gate 4 chặn mọi câu trả lời trích số của tài liệu — kho có dữ liệu mà nói
     "không tìm thấy".

  2. Chế độ khai CHÉP được canh lỏng hơn `facts_ref` rất nhiều: `facts_ref` giải không
     được là lỗi cứng, còn `facts` chép chỉ bị kiểm `src` khác rỗng. Giá trị do LLM gõ
     trở thành sự thật mà Gate 4 dùng để duyệt câu trả lời — Gate 4 đi xác thực LLM
     bằng chính lời khai trước đó của LLM.

Bộ test khoá cả hai, cộng LUẬT OCR (phải chặn ở CODE, không chỉ ở prompt) và việc che
định danh phải theo VỊ TRÍ chứ không theo khoảng cách.

  python3 scripts/van_selftest.py
"""
from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

import ingest_van  # noqa: E402
import lint  # noqa: E402
import numeric_guard  # noqa: E402
import structure  # noqa: E402


# Văn bản nguồn cố ý có BỐN cái bẫy:
#   · 8 ký tự  ngay sau "Điều 7"      -> bẫy che-theo-khoảng-cách
#   · 3 tháng  và 12 tháng            -> hai số cùng đơn vị, khác mục: bẫy gán nhầm
#   · 30 phút  ở mục khác             -> bẫy cửa sổ mục ăn lan
#   · Luật 86/2015/QH13               -> định danh, không được coi là số đo
DOC_PARAGRAPHS = [
    "Chính sách an toàn thông tin",
    "Ban hành theo Luật 86/2015/QH13.",
    "Điều 7. Mật khẩu: độ dài tối thiểu 8 ký tự, đổi định kỳ 3 tháng.",
    "Điều 12. Nhật ký hệ thống được lưu 12 tháng kể từ ngày ghi nhận.",
    "Điều 15. Sự cố phải được báo cáo trong 30 phút.",
    # Số đo DÍNH SÁT locator: bẫy che-theo-khoảng-cách. Cửa sổ của cổng khai báo luôn
    # mở từ chính locator, nên nếu che theo ±12 ký tự thì số 5 này biến mất và một khai
    # báo hoàn toàn đúng bị báo là "không có mặt tại Điều 20".
    "Điều 20: 5 phút là hạn khoá phiên khi không thao tác.",
]

RAW_TEMPLATE = """---
raw_id: {doc_id}
doc_id: {doc_id}
version: 1
kind: van
source_file: originals/{doc_id}.docx
sha256: {sha}
extractor: python-docx
lang: vi
page_type: source
{ocr}generated_by: scripts/extract_van.py
---

{body}
"""

REGISTRY = """documents:
  - doc_id: {doc_id}
    version: 1
    original: originals/{doc_id}.docx
    source_name: {doc_id}.docx
    kind: docx
    sha256: {sha}
    status: canonical
    current: true
    supersedes: null
    visibility: internal
    extractor: van
    raw_paths:
      - raw/{doc_id}.md
"""

DOC_ID = "chinh-sach-attt"


def write_docx(path: Path) -> None:
    import docx

    document = docx.Document()
    for text in DOC_PARAGRAPHS:
        document.add_paragraph(text)
    document.save(path)


def build_corpus(root: Path, *, ocr: bool = False) -> tuple[Path, str]:
    """Corpus tối thiểu, dựng từ .docx THẬT qua đúng extractor của Stage 2.

    Không hardcode nội dung `raw/`: registry đòi `original` tồn tại và sha256 khớp, nên
    fixture đi trọn đường docx → extract_van → raw. Nhờ vậy cổng khai báo được kiểm
    trên văn bản mà extractor thật sinh ra, không phải trên một bản chép tay lý tưởng."""
    import extract_van

    (root / "raw").mkdir(parents=True, exist_ok=True)
    (root / "originals").mkdir(parents=True, exist_ok=True)
    (root / "wiki/sources").mkdir(parents=True, exist_ok=True)
    original = root / f"originals/{DOC_ID}.docx"
    write_docx(original)
    body, _ = extract_van.extract_docx(original)
    sha = extract_van.sha256(original)
    (root / "documents.yml").write_text(
        REGISTRY.format(doc_id=DOC_ID, sha=sha), encoding="utf-8")
    (root / f"raw/{DOC_ID}.md").write_text(
        RAW_TEMPLATE.format(doc_id=DOC_ID, sha=sha, body=body,
                            ocr="ocr: true\n" if ocr else ""),
        encoding="utf-8")
    numeric_guard.reset(root)
    return root, body


def page(declarations: str = "", body: str = "") -> str:
    return f"""---
page: source
name: "Chính sách an toàn thông tin"
doc_id: {DOC_ID}
version: 1
domain: nexus
visibility: internal
raw_paths:
  - raw/{DOC_ID}.md
{declarations}---

# Chính sách an toàn thông tin

{body or "Tài liệu quy định mật khẩu, nhật ký hệ thống và báo cáo sự cố."}

## Nguồn

- `doc_id`: {DOC_ID}
"""


GOOD = (
    f'do_dai_mat_khau_toi_thieu: {{ facts: 8, unit: "ký tự", src: "raw/{DOC_ID}.md :: Điều 7" }}\n'
    f'chu_ky_doi_mat_khau: {{ facts: 3, unit: "tháng", src: "raw/{DOC_ID}.md :: Điều 7" }}\n'
    f'thoi_han_luu_log: {{ facts: 12, unit: "tháng", src: "raw/{DOC_ID}.md :: Điều 12" }}\n'
    f'han_khoa_phien: {{ facts: 5, unit: "phút", src: "raw/{DOC_ID}.md :: Điều 20" }}\n'
)


def check(label: str, condition: bool, detail: str = "") -> bool:
    print(f"  {'✓' if condition else '✗'} {label}" + ("" if condition else f"\n      {detail}"))
    return condition


# --------------------------------------------------------------- Stage 2
def test_extract() -> int:
    """extract_van đọc .docx/.pdf thật: giữ verbatim, giữ thứ tự, phát hiện PDF ảnh."""
    import docx
    import fitz

    failures = 0
    with tempfile.TemporaryDirectory(prefix="pk-van-extract-") as temp:
        temp_dir = Path(temp)

        document = docx.Document()
        document.add_heading("Chính sách an toàn thông tin", level=1)
        document.add_paragraph("Điều 7. Mật khẩu: độ dài tối thiểu 8 ký tự.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Hạng mục"
        table.cell(0, 1).text = "Ngưỡng"
        table.cell(1, 0).text = "Thời hạn lưu log"
        table.cell(1, 1).text = "12 tháng"
        document.add_paragraph("Điều 15. Báo cáo sự cố trong 30 phút.")
        docx_path = temp_dir / "policy.docx"
        document.save(docx_path)

        body, meta = ingest_van_extract_docx(docx_path)
        failures += not check("docx: giữ nguyên văn số đo",
                              "8 ký tự" in body and "12 tháng" in body and "30 phút" in body, body)
        failures += not check("docx: bảng → markdown", "| Thời hạn lưu log | 12 tháng |" in body, body)
        failures += not check("docx: giữ thứ tự đoạn/bảng",
                              body.index("8 ký tự") < body.index("| Hạng mục")
                              < body.index("30 phút"), body)
        failures += not check("docx: đếm đúng đoạn và bảng",
                              meta["tables"] == 1 and meta["paragraphs"] >= 3, str(meta))

        text_pdf = fitz.open()
        text_page = text_pdf.new_page()
        text_page.insert_text((72, 72), "Dieu 7. Mat khau toi thieu 8 ky tu. " * 6)
        pdf_path = temp_dir / "policy.pdf"
        text_pdf.save(pdf_path)
        pdf_body, pdf_meta = ingest_van_extract_pdf(pdf_path)
        failures += not check("pdf có text: chèn marker trang",
                              "[[page 1]]" in pdf_body and pdf_meta["pages"] == 1, pdf_body[:80])

        scan_pdf = fitz.open()
        scan_pdf.new_page()          # trang trắng = 0 ký tự/trang
        scan_path = temp_dir / "scan.pdf"
        scan_pdf.save(scan_path)
        try:
            ingest_van_extract_pdf(scan_path)
            failures += not check("pdf ảnh: đòi OCR thay vì ghi raw rỗng", False,
                                  "không raise NeedOCR")
        except Exception as exc:                       # noqa: BLE001 - kiểm đúng loại
            failures += not check("pdf ảnh: đòi OCR thay vì ghi raw rỗng",
                                  type(exc).__name__ == "NeedOCR", repr(exc))
    return failures


def ingest_van_extract_docx(path):
    import extract_van
    return extract_van.extract_docx(path)


def ingest_van_extract_pdf(path):
    import extract_van
    return extract_van.extract_pdf(path)


# ---------------------------------------------------------------- che định danh
def test_masking() -> int:
    """Định danh phải được che theo VỊ TRÍ, không nuốt số đo đứng cạnh."""
    cases = [
        ("Điều 7: 8 ký tự.", "8", "char"),
        ("Sprint 1 có 20 task.", "20", "task"),
        ("Ô H14 ghi 43 giờ.", "43", "hour"),
    ]
    failures = 0
    for text, want_value, want_unit in cases:
        rows = numeric_guard.transform_numbers(text)
        failures += not check(f"che theo vị trí: {text!r} vẫn thấy {want_value} {want_unit}",
                              any(value == want_value and unit == want_unit
                                  for value, unit, _ in rows), str(rows))
    # Số hiệu mục MỞ ĐẦU DÒNG, không dấu chấm cuối — trang wiki theo chương liệt kê
    # từng Điều nên viết dạng này là tự nhiên. Che theo VỊ TRÍ, không theo hình dạng:
    # số đo nằm TRONG câu, nên '45.5 giờ' giữa câu vẫn phải được soi.
    for text in ("**2.1** Người lao động phải tuân thủ.", "2.1 Nội dung",
                 "- 43.7 Cấm gây gổ.", "### 1.2. Giải thích từ ngữ",
                 # Số thứ tự DANH SÁCH markdown — không có dấu chấm bên trong.
                 "2. **Bản thân Người lao động bị ốm**, thoả mãn điều kiện",
                 "4. Các trường hợp khác được cấp có thẩm quyền chấp thuận",
                 "10) Trường hợp bất khả kháng"):
        rows = numeric_guard.transform_numbers(text)
        failures += not check(f"số hiệu mục đầu dòng được che: {text[:26]!r}",
                              rows == [], str(rows))
    for text, want in [("Tổng cộng 45.5 giờ làm việc.", ("45.5", "hour")),
                       ("40 giờ mỗi tuần.", ("40", "hour")),
                       # Có dấu chấm như số thứ tự, nhưng ngay sau là ĐƠN VỊ -> số đo.
                       ("12. ngày phép", ("12", "day"))]:
        rows = numeric_guard.transform_numbers(text)
        failures += not check(f"số đo vẫn được soi: {text[:26]!r}",
                              [(v, u) for v, u, _ in rows] == [want], str(rows))

    # Văn bản pháp quy tiếng Việt viết tiêu đề chương IN HOA TOÀN BỘ.
    for text in ("CHƯƠNG 2. HỢP ĐỒNG LAO ĐỘNG", "Chương 2. Hợp đồng", "ĐIỀU 45. Kỷ luật"):
        rows = numeric_guard.transform_numbers(text)
        failures += not check(f"số chương/điều được che bất kể hoa thường: {text!r}",
                              rows == [], str(rows))

    # Số hiệu tiêu chuẩn có khoảng trắng — mẫu ô Excel đòi chữ dính số nên bỏ sót.
    # Một trang nhắc ISO 9001 hai lần đủ để pha đếm trị số báo "số mới 9001".
    for text in ("viện dẫn ISO 9001", "theo TCVN 5687:2010", "QCVN 06"):
        rows = numeric_guard.transform_numbers(text)
        failures += not check(f"số hiệu tiêu chuẩn không phải số đo: {text!r}",
                              rows == [], str(rows))

    identifiers = numeric_guard.transform_numbers("Ban hành theo Luật 86/2015/QH13.")
    failures += not check("số hiệu văn bản không bị coi là số đo",
                          all(unit is None for _, unit, _ in identifiers), str(identifiers))

    # OCR đọc dấu chấm cuối số thứ tự mục thành dấu phẩy. Che được, nhưng KHÔNG
    # được nuốt lây dấu thập phân và dấu phân cách nghìn viết theo kiểu Việt Nam.
    comma_clause = numeric_guard.transform_numbers("Khoản 42.29, quy định abc.")
    failures += not check("số thứ tự mục hỏng dấu '42.29,' vẫn được che",
                          comma_clause == [], str(comma_clause))
    for text, want in [("Đạt 43,5 giờ.", "43.5"), ("Phạt 1.234, ghi rõ.", "1234")]:
        rows = numeric_guard.transform_numbers(text)
        failures += not check(f"che dấu phẩy không nuốt {text!r}",
                              any(value == want for value, _, _ in rows), str(rows))

    # Footer chạy trang: số sau nhãn kiểm soát tài liệu là ĐỊNH DANH phiên bản.
    footer = numeric_guard.transform_numbers("Lần ban hành: 1.0\n\nNgày ban hành:\n")
    failures += not check("footer 'Lần ban hành: 1.0' không phải số đo",
                          footer == [], str(footer))

    # Đơn vị phải CÙNG DÒNG với con số. Vắt qua dòng là nhặt nhãn của đoạn sau.
    across = numeric_guard.transform_numbers("Tổng cộng 20\n\nNgày làm việc\n")
    failures += not check("đơn vị không được đọc vắt dòng",
                          across == [("20", None, "20")], str(across))
    inline = numeric_guard.transform_numbers("Nghỉ 12 ngày phép.")
    failures += not check("đơn vị cùng dòng vẫn nhận diện",
                          inline == [("12", "day", "12")], str(inline))
    return failures


# --------------------------------------- Stage 4 · cắt trang theo chương
def test_split_sections() -> int:
    """Tài liệu dài phải sinh thêm trang theo CHƯƠNG, phủ 100% và không mất chữ.

    Trang `source` cố ý là bản tóm tắt, nên với tài liệu 38 trang nó chỉ còn 9% độ
    dài structured. Truy hồi chỉ phục vụ `wiki/`, nên 91% nội dung không bao giờ
    được đánh chỉ mục: kho tìm đúng trang, trang không chứa câu trả lời, và Gate 4
    nói `not_in_kb` cho một tài liệu kho ĐANG CÓ.
    """
    failures = 0
    # Định dạng tiêu đề KHÔNG nhất quán — Stage 3 chạy nhiều khúc độc lập nên chương
    # này có `## `, chương kia là dòng trần. Và mục lục khớp y hệt, chỉ khác số trang.
    front = "Bìa, bảng kiểm soát tài liệu và mục lục. " * 12   # > FRONT_MIN_CHARS
    body = (
        front + "\n\n"
        "CHƯƠNG 1. NHỮNG QUY ĐỊNH CHUNG 5\n"
        "CHƯƠNG 2. HỢP ĐỒNG LAO ĐỘNG 7\n\n"
        "## CHƯƠNG 1. NHỮNG QUY ĐỊNH CHUNG\n\nĐiều 1. Giải thích từ ngữ.\n\n"
        "CHƯƠNG 2. HỢP ĐỒNG LAO ĐỘNG\n\nĐiều 4. Nghỉ 12 ngày phép.\n"
    )
    sections = ingest_van.split_sections(body)
    slugs = [slug for slug, _, _ in sections]
    failures += not check("cắt chương: bỏ dòng mục lục, giữ tiêu đề thân bài",
                          slugs == ["phan-dau", "chuong-1-nhung-quy-dinh-chung",
                                    "chuong-2-hop-dong-lao-dong"], str(slugs))
    failures += not check("cắt chương: nối lại khớp nguyên văn",
                          "".join(text for _, _, text in sections) == body,
                          f"{sum(len(t) for _, _, t in sections)} vs {len(body)}")
    failures += not check("cắt chương: phần bìa/mục lục không bị bỏ",
                          "bảng kiểm soát tài liệu" in sections[0][2], sections[0][2][:40])
    single = ingest_van.split_sections("## CHƯƠNG 1. A\n\nnội dung\n")
    failures += not check("cắt chương: một chương thì không cắt", single == [], str(single))

    # `structured/` mở đầu bằng frontmatter YAML do script sinh. Nếu chương 1 bắt đầu
    # ngay sau đó thì "phần đầu" chỉ còn metadata — cấp một mục rỗng cho Stage 4 thì
    # Opus viết ghi chú "không có nội dung trong bản trích cấp cho Stage 4", và chữ
    # "Stage 4" thành số bịa làm cổng chặn cả tài liệu.
    thin = ("---\ndoc_id: x\nversion: 1\n---\n\n"
            "## CHƯƠNG 1. A\n\nnội dung A\n\n## CHƯƠNG 2. B\n\nnội dung B\n")
    thin_slugs = [slug for slug, _, _ in ingest_van.split_sections(thin)]
    failures += not check("cắt chương: phần đầu chỉ có metadata thì KHÔNG thành mục",
                          "phan-dau" not in thin_slugs, str(thin_slugs))
    return failures


# ------------------------------------------------- Stage 1 · intake khai raw_paths
def test_intake_registers_raw_path() -> int:
    """Extractor sinh đúng một raw/<doc_id>.md thì intake PHẢI khai nó vào registry.

    `van` từng bị bỏ khỏi tập này, nên mọi tài liệu .docx/.pdf nhận `raw_paths: []`
    và Gate 3a chặn 100% với 'tham chiếu raw không thuộc current registry'. Không
    một tài liệu văn xuôi nào có thể vào kho — bằng chứng lane này chưa từng chạy trọn.
    """
    import intake

    failures = 0
    for kind, want in [("docx", "van"), ("pdf", "van"),
                       ("markdown", "markdown"), ("xlsx", "spreadsheet")]:
        got = intake.extractor_for("bat-ky", kind)
        failures += not check(f"intake định tuyến {kind} → {want}", got == want, got)
    single = {"markdown", "spreadsheet", "van"}
    source = Path(intake.__file__).read_text(encoding="utf-8")
    declared = re.search(r'"raw_paths": \[f"raw/\{doc_id\}\.md"\]\s*'
                         r'if extractor in \{([^}]*)\}', source)
    names = set(re.findall(r'"(\w+)"', declared.group(1))) if declared else set()
    failures += not check("intake khai raw_paths cho MỌI lane một-artifact",
                          names == single, f"{sorted(names)} != {sorted(single)}")
    return failures


# ------------------------------------------- Stage 2 · đọc ảnh hai lượt
def test_vision_agreement() -> int:
    """Điều kiện nhận một trang đọc bằng thị giác: hai lượt khớp TỪNG CHỮ SỐ.

    Khác câu chữ thì bỏ qua (xuống dòng, khoảng trắng, cách viết hoa) vì đó là
    cách trình bày. Khác chữ số thì chắc chắn có một lượt sai, và trang đó phải
    được người đối chiếu bản gốc.
    """
    import extract_van

    failures = 0
    same = [
        ("khác cách xuống dòng", "45.4. Nghỉ 365 ngày\ntính từ", "45.4. Nghỉ 365 ngày tính từ"),
        ("khác khoảng trắng", "Làm  40  giờ/tuần", "Làm 40 giờ/tuần"),
    ]
    for label, first, second in same:
        failures += not check(f"hai lượt coi là KHỚP — {label}",
                              extract_van._digits_of(first) == extract_van._digits_of(second),
                              f"{extract_van._digits_of(first)} vs {extract_van._digits_of(second)}")
    differ = [
        ("lệch một chữ số", "thời hạn 365 ngày", "thời hạn 385 ngày"),
        ("thiếu hẳn một số", "Nghỉ 12 ngày phép", "Nghỉ ngày phép"),
    ]
    for label, first, second in differ:
        failures += not check(f"hai lượt coi là LỆCH — {label}",
                              extract_van._digits_of(first) != extract_van._digits_of(second),
                              extract_van._digits_of(first))
    return failures


# ------------------------------------------------------- Stage 3 · cắt khúc
def test_split_chunks() -> int:
    """Cắt khúc phải KHÔNG MẤT ký tự nào và chỉ cắt ở ranh giới trang.

    Đây là bất biến duy nhất giữ cho việc cắt khúc an toàn: nếu nối lại không
    bằng đúng bản gốc thì cổng số của từng khúc vẫn xanh mà cả tài liệu đã hụt.
    """
    body = "".join(f"[[page {n}]]\nĐiều {n}. Nghỉ {n} ngày.\n" + "x" * 4000 + "\n"
                   for n in range(1, 11))
    chunks = structure.split_chunks(body, limit=8000)
    failures = 0
    failures += not check("cắt khúc: nối lại khớp nguyên văn",
                          "".join(chunks) == body, f"{len(''.join(chunks))} vs {len(body)}")
    failures += not check("cắt khúc: tài liệu dài bị chia nhỏ thật", len(chunks) > 1, str(len(chunks)))
    failures += not check("cắt khúc: mỗi khúc mở đầu bằng marker trang",
                          all(chunk.lstrip().startswith("[[page ") for chunk in chunks),
                          str([chunk[:20] for chunk in chunks]))
    short = structure.split_chunks("Chỉ một dòng, 5 ngày.")
    failures += not check("cắt khúc: tài liệu ngắn giữ nguyên một khúc",
                          short == ["Chỉ một dòng, 5 ngày."], str(short))
    return failures


# ------------------------------------------------ cổng biến đổi · hai pha
def test_transform_two_phase() -> int:
    """Pha TRỊ SỐ chống bịa; pha ĐƠN VỊ chỉ xét trị số còn ở cả hai vế.

    Bản trước so theo cặp (trị số, đơn vị) nên hai thao tác dàn lại hợp lệ —
    dàn số vào bảng, và phục hồi dấu tiếng Việt trên từ đơn vị — đều bị báo là
    bịa số. Khoá cả hai chiều: nới đúng hai ca đó, không nới gì khác.
    """
    hard = [
        ("bịa số", "Thời hạn 385 ngày.", "Thời hạn 365 ngày.", "365"),
        ("làm tròn", "Tỉ lệ 45.5 giờ.", "Tỉ lệ 46 giờ.", "46"),
        ("đổi đơn vị thật", "Nghỉ 30 phút.", "Nghỉ 30 giờ.", "đổi đơn vị"),
        ("rơi số có đơn vị", "Nghỉ 45 ngày báo trước.", "Nghỉ báo trước.", "45"),
    ]
    failures = 0
    for label, before, after, needle in hard:
        errors, _ = numeric_guard.check_transform(before, after)
        failures += not check(f"vẫn chặn — {label}",
                              any(needle in message for message in errors), str(errors))
    soft = [
        ("phục hồi dấu trên từ đơn vị", "Làm 40 gid mỗi tuan.", "Làm 40 giờ mỗi tuần."),
        ("dàn số vào bảng", "Làm 40 giờ mỗi tuần.", "| Số | Đơn vị |\n| 40 | giờ |"),
    ]
    for label, before, after in soft:
        errors, warnings = numeric_guard.check_transform(before, after)
        failures += not check(f"chỉ cảnh báo — {label}",
                              errors == [] and warnings != [], f"{errors} / {warnings}")

    # Định danh của vế NGUỒN cũng là định danh ở vế ĐÍCH. `37.4.` trong nguồn được che
    # (số thứ tự mục có chấm cuối); bản tóm tắt nhắc lại thành `37.4` giữa câu và vế
    # kia không che, nên cổng đọc ra "số mới 37.4" — bất đối xứng, không phải bịa số.
    source = ("Điều 37. Mức bồi thường.\n37.4. Khấu trừ tối đa 30% lương.\n"
              "37.5. Trường hợp khác.\nLàm 40 giờ/tuần.")
    quoted = "Điều 37 xen kẽ 37.4–37.5. Làm 40 giờ/tuần."
    errors, _ = numeric_guard.check_transform(source, quoted, allow_loss=True)
    failures += not check("tham chiếu khoản nhắc lại không thành số mới",
                          errors == [], str(errors))
    invented = "Điều 37 xen kẽ 37.4–37.5. Làm 45 giờ/tuần."
    errors, _ = numeric_guard.check_transform(source, invented, allow_loss=True)
    failures += not check("mang định danh sang KHÔNG mở đường cho số đo bịa",
                          any("45" in message for message in errors), str(errors))

    # Cùng một giá trị vừa xuất hiện có đơn vị vừa xuất hiện trần. Rơi bản TRẦN thì
    # bản có đơn vị vẫn còn — không được gán nhãn đơn vị của instance khác rồi nâng
    # thành lỗi cứng.
    before = "Hợp đồng 36 tháng theo mẫu 36 của Bộ. Phụ lục 36 kèm theo."
    after = "Hợp đồng 36 tháng theo mẫu 36 của Bộ."
    errors, warnings = numeric_guard.check_transform(before, after)
    failures += not check("rơi số TRẦN không bị gán đơn vị của instance khác",
                          errors == [] and any("36" in message for message in warnings),
                          f"{errors} / {warnings}")
    dropped = numeric_guard.check_transform("Nghỉ 36 tháng.", "Nghỉ theo quy định.")[0]
    failures += not check("rơi số CÓ đơn vị vẫn là lỗi cứng",
                          any("36" in message and "month" in message for message in dropped),
                          str(dropped))
    return failures


# ------------------------------------------------- Stage 3 · marker provenance
def test_page_markers() -> int:
    """`[[page N]]` không phải số đo (che), nhưng mất nó là mất dấu vết nguồn."""
    failures = 0
    rows = numeric_guard.transform_numbers("[[page 38]]\nNghỉ 12 ngày phép.")
    failures += not check("marker trang không bị coi là số đo",
                          rows == [("12", "day", "12")], str(rows))
    lost = structure.missing_page_markers("[[page 12]]\nA\n[[page 13]]\nB", "[[page 12]]\nA B")
    failures += not check("mất marker trang bị báo là lỗi provenance",
                          lost == ["mất marker trang: [[page 13]]"], str(lost))
    kept = structure.missing_page_markers("[[page 12]]\nA", "# Tiêu đề\n\n[[page 12]]\n\nA")
    failures += not check("giữ đủ marker thì không báo", kept == [], str(kept))

    # Khúc toàn văn xuôi có thể mất nửa nội dung mà không rơi con số nào.
    long_text = "Người lao động có nghĩa vụ tuân thủ nội quy. " * 40
    failures += not check("rút gọn mạnh bị chặn dù không rơi số nào",
                          structure.too_short(long_text, long_text[:len(long_text) // 4]) != [],
                          "không báo")
    failures += not check("dàn lại bình thường không bị chặn",
                          structure.too_short(long_text, "## Nghĩa vụ\n\n" + long_text) == [],
                          "báo nhầm")
    return failures


# --------------------------------------------------------------- Stage 4
def test_stage4_accepts_correct_page(root: Path, structured: str) -> int:
    """Hồi quy chính: trang đúng hợp đồng phải QUA. Trước bản vá thì nó bị chặn."""
    good = page(GOOD, "Mật khẩu tối thiểu 8 ký tự và đổi mỗi 3 tháng theo Điều 7. "
                      "Nhật ký lưu 12 tháng theo Điều 12.")
    problems = ingest_van.validate_page(structured, good, root)
    return not check("trang khai số ĐÚNG hợp đồng đi qua cổng Stage 4",
                     problems == [], "; ".join(problems))


def test_stage4_rejects(root: Path, structured: str) -> int:
    """Mỗi biến thể sai phải bị chặn, và chặn vì ĐÚNG lý do."""
    cases = [
        ("gán nhầm mục: 12 tháng là thời hạn lưu log, không phải chu kỳ đổi mật khẩu",
         f'chu_ky_doi_mat_khau: {{ facts: 12, unit: "tháng", src: "raw/{DOC_ID}.md :: Điều 7" }}\n',
         "không có mặt tại"),
        ("mục không tồn tại trong raw",
         f'chu_ky_doi_mat_khau: {{ facts: 3, unit: "tháng", src: "raw/{DOC_ID}.md :: Điều 99" }}\n',
         "trỏ tới mục không tồn tại"),
        ("đường dẫn raw không có trong corpus",
         'chu_ky_doi_mat_khau: { facts: 3, unit: "tháng", src: "raw/khong-co.md :: Điều 7" }\n',
         "không trỏ tới nguồn raw đọc được"),
        ("src thoát khỏi corpus boundary",
         'chu_ky_doi_mat_khau: { facts: 3, unit: "tháng", src: "../../../etc/passwd :: Điều 7" }\n',
         "không trỏ tới nguồn raw đọc được"),
        ("thiếu unit",
         f'chu_ky_doi_mat_khau: {{ facts: 3, src: "raw/{DOC_ID}.md :: Điều 7" }}\n',
         "thiếu `unit`"),
        ("thiếu src",
         'chu_ky_doi_mat_khau: { facts: 3, unit: "tháng" }\n',
         "thiếu `src`"),
        ("đơn vị lệch so với nguồn",
         f'chu_ky_doi_mat_khau: {{ facts: 3, unit: "giờ", src: "raw/{DOC_ID}.md :: Điều 7" }}\n',
         "đơn vị lệch"),
        ("số hiệu văn bản không phải số đo",
         f'so_hieu: {{ facts: 86, unit: "văn bản", src: "raw/{DOC_ID}.md :: Luật" }}\n',
         "không có mặt tại"),
        ("facts không phải số",
         f'chu_ky_doi_mat_khau: {{ facts: "ba tháng", unit: "tháng", src: "raw/{DOC_ID}.md :: Điều 7" }}\n',
         "không phải số đo hay ngày"),
    ]
    failures = 0
    for label, declaration, want in cases:
        problems = ingest_van.validate_page(structured, page(declaration), root)
        failures += not check(f"chặn — {label}",
                              any(want in message for message in problems),
                              f"mong khớp {want!r}, nhận: {problems}")

    invented = page(GOOD, "Mật khẩu tối thiểu 8 ký tự; hệ thống chịu tải 500 phiên.")
    problems = ingest_van.validate_page(structured, invented, root)
    failures += not check("chặn — thân bài bịa số không có trong nguồn",
                          any("500" in message for message in problems), str(problems))

    rounded = page(GOOD, "Sự cố phải báo cáo trong 0.5 phút.")
    problems = ingest_van.validate_page(structured, rounded, root)
    failures += not check("chặn — thân bài đổi/làm tròn số",
                          problems != [], str(problems))
    return failures


# --------------------------------------------------------------- LUẬT OCR
def test_ocr_rule() -> int:
    """Trang dựng từ nguồn OCR không được khai số — kể cả khi trang KHÔNG tự khai `ocr`.

    Trạng thái OCR phải lấy từ raw (script ghi), không từ frontmatter trang (LLM ghi):
    prompt bảo LLM đừng khai facts, nhưng lời dặn không phải là cổng."""
    failures = 0
    with tempfile.TemporaryDirectory(prefix="pk-van-ocr-") as temp:
        root, structured = build_corpus(Path(temp), ocr=True)
        problems = ingest_van.validate_page(structured, page(GOOD), root)
        failures += not check("LUẬT OCR: chặn khai số từ nguồn OCR",
                              any("LUẬT OCR" in message for message in problems), str(problems))

        clean = ingest_van.validate_page(structured, page(""), root)
        failures += not check("LUẬT OCR: trang OCR không khai số vẫn đi qua",
                              clean == [], str(clean))
    return failures


# --------------------------------------------------------------- Gate 3a
def test_gate3a(root: Path) -> int:
    """lint-numbers phải đỏ với khai báo chép sai, ngang mức với facts_ref hỏng."""
    import yaml

    def run(declaration: str):
        frontmatter, _ = numeric_guard.split_frontmatter(page(declaration))
        lint.errors.clear()
        lint.warns.clear()
        lint.lint_numbers({"wiki/sources/x.md": {"fm": frontmatter, "body": ""}}, {}, root)
        return [message for _, _, message in lint.errors]

    assert yaml  # frontmatter đã được parse bằng yaml qua split_frontmatter
    failures = 0
    failures += not check("Gate 3a xanh với khai báo đúng", run(GOOD) == [], str(run(GOOD)))
    bad = run(f'chu_ky_doi_mat_khau: {{ facts: 12, unit: "tháng", '
              f'src: "raw/{DOC_ID}.md :: Điều 7" }}\n')
    failures += not check("Gate 3a đỏ với giá trị gán nhầm mục",
                          any("không có mặt tại" in message for message in bad), str(bad))
    return failures


# --------------------------------------------------------------- Gate 4
def test_gate4_distrusts_unverified(root: Path) -> int:
    """Runtime không được đăng ký một khai báo chưa đối chiếu được.

    Gate 3a đã chặn lúc xuất bản; đây là lớp fail-closed thứ hai, để một trang lọt vào
    kho bằng đường khác cũng không biến con số của LLM thành sự thật."""
    wiki = root / f"wiki/sources/{DOC_ID}.md"
    failures = 0

    wiki.write_text(page(GOOD), encoding="utf-8")
    numeric_guard.reset(root)
    guard = numeric_guard.AnswerGuard(root)
    cite = [f"wiki/sources/{DOC_ID}.md"]
    failures += not check("Gate 4: số khai ĐÚNG được đăng ký và trả lời được",
                          guard.check("Mật khẩu tối thiểu 8 ký tự.", cites=cite) == [],
                          str(guard.check("Mật khẩu tối thiểu 8 ký tự.", cites=cite)))

    wiki.write_text(page(
        f'chu_ky_doi_mat_khau: {{ facts: 7, unit: "tháng", '
        f'src: "raw/{DOC_ID}.md :: Điều 7" }}\n'), encoding="utf-8")
    numeric_guard.reset(root)
    guard = numeric_guard.AnswerGuard(root)
    failures += not check("Gate 4: khai báo không đối chiếu được KHÔNG mở khoá số",
                          guard.check("Đổi mật khẩu mỗi 7 tháng.", cites=cite) == ["7"],
                          str(guard.check("Đổi mật khẩu mỗi 7 tháng.", cites=cite)))
    wiki.unlink()
    numeric_guard.reset(root)
    return failures


def main() -> int:
    failures = 0
    print("── Stage 2 · EXTRACT (.docx / .pdf thật) ──")
    failures += test_extract()
    print("── numeric_guard · che định danh theo vị trí ──")
    failures += test_masking()
    print("── Stage 4 · cắt trang theo chương ──")
    failures += test_split_sections()
    print("── Stage 1 · intake khai raw_paths ──")
    failures += test_intake_registers_raw_path()
    print("── Stage 2 · đọc ảnh hai lượt đối chiếu ──")
    failures += test_vision_agreement()
    print("── Stage 3 · cắt khúc không mất chữ ──")
    failures += test_split_chunks()
    print("── numeric_guard · cổng biến đổi hai pha ──")
    failures += test_transform_two_phase()
    print("── Stage 3 · marker trang là provenance ──")
    failures += test_page_markers()
    with tempfile.TemporaryDirectory(prefix="pk-van-") as temp:
        root, structured = build_corpus(Path(temp))
        print("── Stage 4 · cổng ghi trang wiki ──")
        failures += test_stage4_accepts_correct_page(root, structured)
        failures += test_stage4_rejects(root, structured)
        print("── LUẬT OCR ──")
        failures += test_ocr_rule()
        print("── Gate 3a · lint-numbers ──")
        failures += test_gate3a(root)
        print("── Gate 4 · runtime fail-closed ──")
        failures += test_gate4_distrusts_unverified(root)
        numeric_guard.reset(root)
    if failures:
        print(f"\n✗ VAN self-test: {failures} kiểm tra FAIL")
        return 1
    print("\n✓ VAN self-test: .docx/.pdf → raw → Stage 4 → Gate 3a → Gate 4")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
